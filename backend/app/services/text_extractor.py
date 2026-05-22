import pdfplumber
from docx import Document
from pathlib import Path
from typing import Optional
import logging

logger = logging.getLogger(__name__)

class TextExtractor:
    """Extract text from PDF and DOCX files"""
    
    @staticmethod
    async def extract_from_pdf(file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            extracted_text = "\n".join(text_parts)
            logger.info(f"Extracted {len(extracted_text)} chars from PDF: {file_path.name}")
            return extracted_text.strip()
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    @staticmethod
    async def extract_from_docx(file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            extracted_text = "\n".join(text_parts)
            logger.info(f"Extracted {len(extracted_text)} chars from DOCX: {file_path.name}")
            return extracted_text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    @staticmethod
    async def extract_text(file_path: Path) -> str:
        """Extract text based on file extension"""
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return await TextExtractor.extract_from_pdf(file_path)
        elif suffix == '.docx':
            return await TextExtractor.extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
